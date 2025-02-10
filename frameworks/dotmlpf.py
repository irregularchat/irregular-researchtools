# /frameworks/DOTMLPF.py
"""
DOTMLPF Analysis Framework
"""
import streamlit as st
from utilities.gpt import chat_gpt
from dotenv import load_dotenv
import json
import re

load_dotenv()

def dotmlpf_page():
    # Initialize session storage for AI responses to ensure they persist.
    if "dotmlpf_summary" not in st.session_state:
        st.session_state["dotmlpf_summary"] = ""
    if "tradoc_alignment" not in st.session_state:
        st.session_state["tradoc_alignment"] = ""
    if "command_endorsement" not in st.session_state:
        st.session_state["command_endorsement"] = ""
    if "improved_problem_statement" not in st.session_state:
        st.session_state["improved_problem_statement"] = ""

    st.title("DOTMLPF-P Analysis Framework")
    st.write("""
    The DOTMLPF-P Analysis Framework helps you assess organizational capabilities across several key areas:
    
    - **Doctrine**: The principles and strategies guiding operations.
    - **Organization**: The structure and distribution of forces.
    - **Training**: The preparedness and skills of personnel.
    - **Materiel**: The equipment, technology, and resources available.
    - **Leadership**: The quality and effectiveness of command.
    - **Personnel**: The manpower, expertise, and morale.
    - **Facilities**: The infrastructure and support systems.
    - **Policy**: The policies and frameworks governing procedures and resource allocation.
    
    You can use this framework to analyze friendly forces, adversary forces, or our own.
    """)

    # Toggle for including Joint Integration category
    st.write("Optional: Include Joint Integration analysis to assess interoperability with partner forces.")
    include_joint_integration = st.checkbox("Include Joint Integration category")

    # Input field for Organization
    organization_input = st.text_input("Enter Organization Details:", max_chars=240)

    # Choose the force type to analyze
    force_type = st.selectbox("Select Force Type", ["Friendly", "Adversary", "Our Own"])

    # New input field for the overall goal of the force
    force_goal_input = st.text_input("Enter Goal of the Force:", max_chars=240)

    # Prompt the user for their analysis goal (distinct from the force goal)
    goal_input = st.text_input("Enter Goal of the Analysis:", max_chars=240)

    # Provide guiding questions for initial problem definition
    st.markdown("**Guiding Questions for Defining Your Problem Statement**")
    st.write("""
    1. What is the primary objective or end state you want to achieve through this analysis?
    2. How does this analysis align with broader strategic or organizational goals?
    3. Which stakeholders are most impacted, and what outcomes are they expecting?
    4. What constraints, timelines, or resources (e.g., budget, manpower) shape your current challenges?
    5. Are there any known threats, gaps, or shortfalls that precipitated this analysis?
    """)

    # If "Our Own" is selected, request details about an operational gap
    if force_type == "Our Own":
        operational_gap_input = st.text_area(
            "Describe any operational gap or capability shortfall preventing your mission requirement from being met:",
            max_chars=1000
        )
    else:
        operational_gap_input = ""

    st.markdown("---")

    # DOTMLPF-P categories
    dotmlpf_categories = [
        "Doctrine",
        "Organization",
        "Training",
        "Materiel",
        "Leadership",
        "Personnel",
        "Facilities",
        "Policy"
    ]

    # If user opts in, add Joint Integration category
    if include_joint_integration:
        dotmlpf_categories.append("Joint Integration")

    # Dictionary to hold user inputs for each category
    user_inputs = {}

    # Loop through each category: display text area & AI suggestion button
    for cat in dotmlpf_categories:
        st.subheader(cat)

        # If we haven't yet stored user analysis for this category, initialize it
        if cat not in st.session_state:
            st.session_state[cat] = ""

        user_text = st.text_area(f"Enter your analysis/observations for {cat}:", key=cat)
        user_inputs[cat] = user_text

        # If we haven't stored any questions for this category, initialize an empty list
        if f'analysis_questions_{cat}' not in st.session_state:
            st.session_state[f'analysis_questions_{cat}'] = []

        # AI suggestion button for generating analysis questions
        if st.button(f"AI: Suggest {cat} Questions", key=f"ai_{cat}"):
            try:
                # Base prompt referencing TRADOC, JCIDS, and CBA for all force types
                base_system_prompt = (
                    "You are an experienced military capability analyst specializing in DOTMLPF-P assessments. "
                    f"Your role is to evaluate the following force type: {force_type}.\n\n"
                    "In accordance with the TRADOC Capability Development Document (CDD) guidelines, your analysis should:\n"
                    "• Identify existing capabilities in each DOTMLPF-P domain.\n"
                    "• Highlight any shortfalls or missing elements in those capabilities.\n"
                    "• Use Capability-Based Assessment (CBA) methodology to understand potential gaps.\n"
                    "• Reference Key Performance Parameters (KPPs) and Key System Attributes (KSAs) for performance tracking.\n"
                    "• Consider System of Systems (SoS) dependencies and operational risks.\n"
                    "• Address Doctrine, Organization, Training, Materiel, Leadership, Personnel, Facilities, and Policy.\n\n"
                    "For the following category, produce three specific, measurable, and actionable questions/prompts focused on:\n"
                    "1) Identifying the existing capabilities.\n"
                    "2) Determining gaps or deficiencies within those capabilities.\n"
                    "3) Noting how these align with TRADOC and—where applicable—JCIDS criteria.\n"
                    "Please provide your response STRICTLY in JSON format as follows:\n"
                    "{\"questions\": [\"Question 1\", \"Question 2\", \"Question 3\"]}\n"
                )

                # Add specialized guidance when force_type == "Our Own"
                if force_type == "Our Own":
                    base_system_prompt += (
                        "\nSince this analysis focuses on 'Our Own' forces, also ensure:\n"
                        "• Any operational gap provided is examined in terms of existing capabilities.\n"
                        "• The user understands what is currently in place (e.g., doctrine, units, training) and what is missing or incomplete.\n"
                        "• Consider potential resource pathways or policy changes if shortfalls are identified.\n"
                    )

                system_msg = {"role": "system", "content": base_system_prompt}

                # Build user message
                if force_type == "Our Own":
                    user_msg_content = (
                        f"Force Type: {force_type}\n"
                        f"Goal of the Force: {force_goal_input}\n"
                        f"Goal of the Analysis: {goal_input}\n"
                        f"Organization: {organization_input}\n"
                        f"Operational Gap: {operational_gap_input}\n"
                        f"Category: {cat}\n"
                        f"Current Input Provided: {user_text}\n\n"
                        "Please generate three specific, measurable, and actionable questions/prompts as described above.\n"
                        "Remember to output your response in the following JSON format exactly:\n"
                        "{\"questions\": [\"Question 1\", \"Question 2\", \"Question 3\"]}"
                    )
                else:
                    user_msg_content = (
                        f"Force Type: {force_type}\n"
                        f"Goal of the Force: {force_goal_input}\n"
                        f"Goal of the Analysis: {goal_input}\n"
                        f"Organization: {organization_input}\n"
                        f"Category: {cat}\n"
                        f"Current Input Provided: {user_text}\n\n"
                        "Please generate three specific, measurable, and actionable questions/prompts as described above.\n"
                        "Remember to output your response in the following JSON format exactly:\n"
                        "{\"questions\": [\"Question 1\", \"Question 2\", \"Question 3\"]}"
                    )

                user_msg = {"role": "user", "content": user_msg_content}
                ai_response = chat_gpt([system_msg, user_msg], model="gpt-4o-mini")

                # Try to parse the AI response as JSON.
                try:
                    data = json.loads(ai_response)
                    question_list = data.get("questions", [])
                    # Ensure we have three questions; if not, fallback.
                    if not isinstance(question_list, list) or len(question_list) < 3:
                        raise ValueError("JSON did not contain at least three questions.")
                except Exception as json_err:
                    # Fallback to regex extraction if JSON parsing fails.
                    lines = ai_response.strip().split('\n')
                    question_list = []
                    for line in lines:
                        line_stripped = line.strip()
                        if re.match(r'^(\d[\.\)]|•|-)\s', line_stripped) or re.match(r'^\d\)', line_stripped):
                            question_list.append(line_stripped)
                    # If we got fewer than 3 ideas, fallback to the entire response.
                    if len(question_list) < 3:
                        question_list = [ai_response]

                # Store parsed questions in session state
                st.session_state[f'analysis_questions_{cat}'] = question_list

            except Exception as e:
                st.error(f"Error generating AI suggestion for {cat}: {e}")

            st.markdown("---")

        # Display previously generated questions with associated answer input fields and answer template option
        if st.session_state[f'analysis_questions_{cat}']:
            st.markdown("**AI-Suggested Questions/Prompts:**")
            for idx, question in enumerate(st.session_state[f'analysis_questions_{cat}']):
                st.write(f"**Q{idx+1}:** {question}")
                # Create an answer key for each question
                answer_key = f"analysis_answer_{cat}_{idx}"
                if answer_key not in st.session_state:
                    st.session_state[answer_key] = ""
                st.session_state[answer_key] = st.text_area(
                    f"Your answer to Q{idx+1} (optional):",
                    st.session_state[answer_key],
                    key=f"{cat}_q{idx}_input"
                )
                # Button to generate a fill-in-the-blank answer template using available variables
                template_key = f"template_answer_{cat}_{idx}"
                if st.button(f"Generate Answer Template for Q{idx+1}", key=f"template_btn_{cat}_{idx}"):
                    try:
                        template_prompt = (
                            f"You are an experienced military capability analyst. A user is performing research on '{cat}' as part of a DOTMLPF-P analysis.\n"
                            f"Context:\n"
                            f"Force Type: {force_type}\n"
                            f"Goal of the Force: {force_goal_input}\n"
                            f"Goal of the Analysis: {goal_input}\n"
                            f"Organization: {organization_input}\n"
                        )
                        if force_type == "Our Own":
                            template_prompt += f"Operational Gap: {operational_gap_input}\n"
                        template_prompt += (
                            f"\nAnalysis Question: {question}\n\n"
                            "Generate a fill-in-the-blank answer template that includes variable placeholders (e.g., [X], [Y], [Z]) "
                            "to guide the user in answering this question. For example, in the Doctrine category, a template might be:\n"
                            "\"For Doctrine, the force has [X] doctrinal frameworks such as [Y]. However, they are limited by [Z] constraints.\"\n\n"
                            "Return just the template answer."
                        )
                        template_response = chat_gpt([{"role": "system", "content": template_prompt}], 
                                                     model="gpt-4o-mini")
                        st.session_state[template_key] = template_response
                    except Exception as e:
                        st.error(f"Error generating template answer for Q{idx+1} in {cat}: {e}")
                # Display the generated answer template if available.
                if template_key in st.session_state and st.session_state[template_key]:
                    st.text_area(f"Answer Template for Q{idx+1}:", st.session_state[template_key],
                                 key=f"{cat}_template_q{idx}_input")

    # Button to generate a consolidated summary from all categories
    if st.button("Generate Consolidated DOTMLPF-P Summary"):
        try:
            summary_prompt = "Based on the following DOTMLPF-P observations, provide a concise summary with key insights and TRADOC-aligned recommendations:\n"
            insufficient_categories = []
            for cat in dotmlpf_categories:
                # Only include category if observations or answers were provided.
                analysis = user_inputs.get(cat, "").strip()
                answers_for_cat = ""
                if st.session_state.get(f'analysis_questions_{cat}', []):
                    for idx, question in enumerate(st.session_state.get(f'analysis_questions_{cat}')):
                        user_answer = st.session_state.get(f'analysis_answer_{cat}_{idx}', "").strip()
                        if user_answer:
                            answers_for_cat += f"\nQ: {question}\nA: {user_answer}\n"
                if analysis or answers_for_cat:
                    summary_prompt += f"\n{cat}:\n{analysis}\n{answers_for_cat}\n"
                else:
                    insufficient_categories.append(cat)
            if insufficient_categories:
                summary_prompt += "\nNote: Insufficient data was provided for the following categories: " + ", ".join(insufficient_categories) + "."
            if force_type == "Our Own":
                summary_prompt += (
                    "\nSince these are 'Our Own' forces, ensure the summary highlights:\n"
                    "• Which capabilities already exist within each domain\n"
                    "• How the operational gap (if any) may reveal shortfalls or deficiencies\n"
                    "• Relevant resourcing or doctrinal references, if the shortfalls require updates\n"
                )

            system_msg = {
                "role": "system",
                "content": (
                    f"You are an expert military strategist focused on the following:\n"
                    f"Goal of the Force: {force_goal_input}\n"
                    f"Goal of the Analysis: {goal_input}\n\n"
                    "Summarize the DOTMLPF-P analysis below in alignment with TRADOC guidelines and any operational gap (if provided). "
                    "Emphasize what existing capabilities each domain offers, what might be missing, "
                    "and the impact on overall capability."
                )
            }

            if force_type == "Our Own":
                user_msg_content = f"Operational Gap: {operational_gap_input}\n\n{summary_prompt}"
            else:
                user_msg_content = summary_prompt

            user_msg = {"role": "user", "content": user_msg_content}
            summary_response = chat_gpt([system_msg, user_msg], model="gpt-4o-mini")

            st.session_state["dotmlpf_summary"] = summary_response

            st.subheader("Consolidated DOTMLPF-P Summary")
            st.write(summary_response)

        except Exception as e:
            st.error(f"Error generating summary: {e}")

    st.markdown("---")
    # New section: Improve Problem Statement after DOTMLPF-P categories.
    st.subheader("Improve Problem Statement")
    # Use a text area that lets the user edit the statement.
    # If left blank, an initial problem statement will be generated.
    current_problem_statement = st.text_area(
        "Enter or edit your problem statement (if left blank, an initial statement will be generated from your data):",
        key="problem_statement",
        height=150
    )
    if st.button("Generate/Improve Problem Statement with AI", key="improve_prob_statement"):
        if current_problem_statement.strip() == "":
            # No initial input provided => generate a new problem statement based on all user data.
            prompt = (
                "You are an expert in crafting clear, actionable, and research-informed problem statements. "
                "Based on the following analysis details, generate an initial problem statement that encapsulates the key challenges, gaps, "
                "and strategic imperatives derived from the DOTMLPF-P analysis. Include considerations of force type, goals, organization, and any operational gaps noted.\n\n"
                f"Force Type: {force_type}\n"
                f"Force Goal: {force_goal_input}\n"
                f"Goal of the Analysis: {goal_input}\n"
                f"Organization: {organization_input}\n"
                f"Operational Gap: {operational_gap_input if force_type == 'Our Own' else 'N/A'}\n\n"
                f"DOTMLPF-P Analysis Summary:\n{st.session_state.get('dotmlpf_summary', 'No summary available.')}\n\n"
                "Return only the generated problem statement."
            )
            generated_statement = chat_gpt([{"role": "system", "content": prompt}], model="gpt-4o-mini")
            # Update the text area with the generated problem statement so that the user can further edit it.
            st.session_state["problem_statement"] = generated_statement
            st.success("Generated Problem Statement:")
            st.write(generated_statement)
        else:
            # A problem statement was provided => improve it.
            prompt = (
                "You are an expert in crafting clear, actionable, and research-informed problem statements. "
                "Based on the following consolidated DOTMLPF-P analysis and the original problem statement provided below, "
                "please improve the problem statement to make it more concise, focused, and reflective of the analysis.\n\n"
                f"DOTMLPF-P Analysis Summary:\n{st.session_state.get('dotmlpf_summary', 'No summary available.')}\n\n"
                f"Original Problem Statement: {current_problem_statement}\n\n"
                "Return only the improved problem statement."
            )
            improved_statement = chat_gpt([{"role": "system", "content": prompt}], model="gpt-4o-mini")
            # Update the text area with the improved statement.
            st.session_state["problem_statement"] = improved_statement
            st.success("Improved Problem Statement:")
            st.write(improved_statement)

    st.markdown("---")
    st.subheader("TRADOC Alignment")
    st.write("Click below to generate TRADOC alignment considerations—covering JCIDS, KPPs/KSAs, and resourcing strategies.")
    if st.button("Generate TRADOC Alignment"):
        try:
            alignment_prompt = (
                "You are an expert in TRADOC capability development and Joint Capabilities Integration Development System (JCIDS). "
                "Given the following DOTMLPF-P Summary:\n\n"
                f"{st.session_state['dotmlpf_summary']}\n\n"
                "Identify specific ways to align findings with:\n"
                "- JCIDS requirements or acquisition milestones\n"
                "- Key Performance Parameters (KPPs) and Key System Attributes (KSAs)\n"
                "- Resourcing options (POM, APFIT, Rapid Prototyping)\n\n"
                "Focus on how current capabilities and shortfalls map to these frameworks. "
                "Structure the output as bullet points or short paragraphs so it can be easily referenced by senior leaders."
            )

            alignment_msg = {
                "role": "system",
                "content": alignment_prompt
            }
            alignment_response = chat_gpt([alignment_msg], model="gpt-4o-mini")
            st.session_state["tradoc_alignment"] = alignment_response

            st.write(alignment_response)

        except Exception as e:
            st.error(f"Error generating TRADOC alignment: {e}")

    st.markdown("---")
    if force_type == "Our Own":
        st.subheader("Command Endorsement Strategy")
        st.write("Propose how leadership at various levels should acknowledge or resource the capabilities/shortfalls identified.")
        if st.button("Generate Command Endorsement Strategy"):
            try:
                endorsement_prompt = (
                    "You are a senior-level advisor with the authority to recommend endorsements at multiple echelons.\n\n"
                    "Given the following DOTMLPF-P summary (including any operational gap and recommendations):\n\n"
                    f"{st.session_state['dotmlpf_summary']}\n\n"
                    "Provide specific command endorsement strategies for:\n"
                    "1. **Operational Level** (e.g., Force providers, TRADOC Centers of Excellence)\n"
                    "2. **Strategic Level** (e.g., Army Futures Command, ASA(ALT), Joint Staff)\n"
                    "3. **Policy Level** (e.g., DoD Directives, Congressional funding requests)\n\n"
                    "For each level, identify the relevant command authorities and detail what kind of endorsement is needed, "
                    "such as:\n"
                    "- A **doctrinal update** (e.g., changes to TRADOC directives, Army Regulations)\n"
                    "- A **resource allocation** (e.g., applying for APFIT, rapid prototyping funds, or POM adjustments)\n"
                    "- A **force structure change** (e.g., reorganizing units or re-aligning personnel)\n"
                    "- A **new acquisition program** (e.g., starting or modifying a Program of Record)\n\n"
                    "Tie each recommendation to the existing capabilities or shortfalls identified in the summary. "
                    "Where applicable, reference Army Warfighting Challenges (AWFCs) or Program Objective Memorandum (POM) "
                    "cycles to illustrate how these endorsements fit into broader Joint Force and TRADOC planning."
                )

                command_msg = {
                    "role": "system",
                    "content": endorsement_prompt
                }
                command_response = chat_gpt([command_msg], model="gpt-4o-mini")

                st.session_state["command_endorsement"] = command_response
                st.write(command_response)

            except Exception as e:
                st.error(f"Error generating command endorsement strategy: {e}")

    st.markdown("---")
    if st.button("Export DOTMLPF-P Analysis as JSON"):
        cat_qa_data = {}
        for cat in dotmlpf_categories:
            cat_qa_data[cat] = {
                "analysis_observations": user_inputs.get(cat, ""),
                "questions_answers": []
            }
            q_list = st.session_state.get(f'analysis_questions_{cat}', [])
            for idx, question in enumerate(q_list):
                user_answer = st.session_state.get(f'analysis_answer_{cat}_{idx}', "")
                cat_qa_data[cat]["questions_answers"].append({
                    "question": question,
                    "answer": user_answer
                })

        analysis_data = {
            "force_type": force_type,
            "force_goal": force_goal_input,
            "goal_of_analysis": goal_input,
            "organization": organization_input,
            "operational_gap": operational_gap_input if force_type == "Our Own" else "",
            "DOTMLPF-P": cat_qa_data,
            "TRADOC_Alignment": st.session_state["tradoc_alignment"],
            "Command_Endorsement": st.session_state["command_endorsement"] if force_type == "Our Own" else ""
        }

        json_data = json.dumps(analysis_data, indent=2)
        file_name = f"DOTMLPF-P_Analysis_{force_type}.json"

        st.download_button(
            label="Download JSON",
            data=json_data,
            file_name=file_name,
            mime="application/json"
        )

    # Existing JSON export button is above; now add export as PDF/DOCX option.
    export_format = st.selectbox("Select export format for document", ["PDF", "DOCX"], key="export_format")
    if st.button("Export DOTMLPF-P Analysis as Document"):
        # Build document title based on organization_input; provide a default if blank.
        title_doc = f"DOTMLPF-P Analysis for {organization_input}" if organization_input.strip() else "DOTMLPF-P Analysis"
        # Retrieve key sections from session state.
        problem_statement = st.session_state.get("problem_statement", "No problem statement provided.")
        dotmlpf_summary = st.session_state.get("dotmlpf_summary", "No summary available.")
        tradoc_alignment = st.session_state.get("tradoc_alignment", "")
        command_endorsement = st.session_state.get("command_endorsement", "")
        
        # Export as DOCX using python-docx.
        if export_format == "DOCX":
            try:
                from docx import Document
                from docx.shared import Pt
                import io

                document = Document()
                # Set document's default style to Arial 12pt.
                style = document.styles["Normal"]
                font = style.font
                font.name = "Arial"
                font.size = Pt(12)

                document.add_heading(title_doc, level=0)
                document.add_heading("Problem Statement", level=1)
                document.add_paragraph(problem_statement)
                document.add_heading("DOTMLPF-P Summary", level=1)
                document.add_paragraph(dotmlpf_summary)
                
                # Loop through DOTMLPF-P categories.
                for cat in dotmlpf_categories:
                    document.add_heading(f"Category: {cat}", level=2)
                    analysis_observations = st.session_state.get(cat, "")
                    if analysis_observations:
                        document.add_paragraph("Observations: " + analysis_observations)
                    # Add questions and answers if available.
                    questions = st.session_state.get(f"analysis_questions_{cat}", [])
                    if questions:
                        document.add_paragraph("Questions and Answers:")
                        for idx, question in enumerate(questions):
                            document.add_paragraph(f"Q: {question}", style="List Bullet")
                            answer = st.session_state.get(f"analysis_answer_{cat}_{idx}", "")
                            document.add_paragraph(f"A: {answer}", style="List Bullet")
                
                if tradoc_alignment:
                    document.add_heading("TRADOC Alignment", level=1)
                    document.add_paragraph(tradoc_alignment)
                if command_endorsement:
                    document.add_heading("Command Endorsement", level=1)
                    document.add_paragraph(command_endorsement)
                
                # Save the document to an in-memory binary stream.
                docx_io = io.BytesIO()
                document.save(docx_io)
                docx_io.seek(0)
                st.download_button(
                    label="Download DOCX",
                    data=docx_io,
                    file_name=f"{title_doc}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Error generating DOCX export: {e}")
        else:  # Export as PDF using fpdf.
            try:
                from fpdf import FPDF
                import io

                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                
                # Title centered.
                pdf.cell(0, 10, txt=title_doc, ln=True, align="C")
                pdf.ln(5)
                pdf.cell(0, 10, txt="Problem Statement", ln=True)
                pdf.multi_cell(0, 10, txt=problem_statement)
                pdf.ln(5)
                pdf.cell(0, 10, txt="DOTMLPF-P Summary", ln=True)
                pdf.multi_cell(0, 10, txt=dotmlpf_summary)
                pdf.ln(5)
                
                for cat in dotmlpf_categories:
                    pdf.cell(0, 10, txt=f"Category: {cat}", ln=True)
                    analysis_observations = st.session_state.get(cat, "")
                    if analysis_observations:
                        pdf.multi_cell(0, 10, txt="Observations: " + analysis_observations)
                    questions = st.session_state.get(f"analysis_questions_{cat}", [])
                    if questions:
                        pdf.cell(0, 10, txt="Questions and Answers:", ln=True)
                        for idx, question in enumerate(questions):
                            answer = st.session_state.get(f"analysis_answer_{cat}_{idx}", "")
                            pdf.multi_cell(0, 10, txt=f"Q: {question}")
                            pdf.multi_cell(0, 10, txt=f"A: {answer}")
                    pdf.ln(5)
                
                if tradoc_alignment:
                    pdf.cell(0, 10, txt="TRADOC Alignment", ln=True)
                    pdf.multi_cell(0, 10, txt=tradoc_alignment)
                    pdf.ln(5)
                if command_endorsement:
                    pdf.cell(0, 10, txt="Command Endorsement", ln=True)
                    pdf.multi_cell(0, 10, txt=command_endorsement)
                
                # Output the PDF to a binary string.
                pdf_output = pdf.output(dest="S").encode("latin1")
                st.download_button(
                    label="Download PDF",
                    data=pdf_output,
                    file_name=f"{title_doc}.pdf",
                    mime="application/pdf"
                )
            except Exception as e:
                st.error(f"Error generating PDF export: {e}")

def main():
    dotmlpf_page()

if __name__ == "__main__":
    main()