import os
import json
import io
import streamlit as st
from typing import Dict, List, Any, Optional, Union, Callable, Tuple
from abc import ABC, abstractmethod
from datetime import datetime

class BaseFramework(ABC):
    """Base class for all analytical frameworks"""
    
    def __init__(self, name: str):
        self.name = name
        self.questions: Dict[str, List[str]] = {}
        self.responses: Dict[str, Any] = {}
        self.components: List[str] = []
        self.metadata: Dict[str, Any] = {
            "created": datetime.now().isoformat(),
            "last_modified": datetime.now().isoformat(),
            "version": "1.0"
        }
    
    @abstractmethod
    def generate_questions(self, component: str) -> List[str]:
        """Generate questions for a specific component"""
        pass
    
    def generate_all_questions(self) -> Dict[str, List[str]]:
        """Generate questions for all components"""
        result = {}
        for component in self.components:
            result[component] = self.generate_questions(component)
        return result
    
    def set_response(self, component: str, response: Any) -> None:
        """Set response for a component"""
        self.responses[component] = response
        self.metadata["last_modified"] = datetime.now().isoformat()
    
    def get_response(self, component: str) -> Any:
        """Get response for a component"""
        return self.responses.get(component)
    
    def get_all_responses(self) -> Dict[str, Any]:
        """Get all responses"""
        return self.responses
    
    def export_to_json(self, filepath: Optional[str] = None) -> str:
        """Export framework data to JSON"""
        data = {
            "framework": self.name,
            "components": self.components,
            "questions": self.questions,
            "responses": self.responses,
            "metadata": self.metadata
        }
        
        if filepath:
            with open(filepath, 'w') as f:
                json.dump(data, f, indent=2)
                
        return json.dumps(data, indent=2)
    
    def import_from_json(self, json_data: str) -> bool:
        """Import framework data from JSON"""
        try:
            data = json.loads(json_data)
            
            # Validate that this is the correct framework
            if data.get("framework") != self.name:
                return False
                
            # Import components, questions, and responses
            self.components = data.get("components", [])
            self.questions = data.get("questions", {})
            self.responses = data.get("responses", {})
            self.metadata = data.get("metadata", self.metadata)
            
            # Update last modified timestamp
            self.metadata["last_modified"] = datetime.now().isoformat()
            
            return True
        except Exception as e:
            print(f"Error importing data: {e}")
            return False
    
    def select_all(self):
        """Generate analysis for all components"""
        return self.generate_all_questions()
    
    def initialize_session_state(self, prefix: str, defaults: Dict[str, Any]) -> None:
        """Initialize session state variables with a prefix to avoid collisions"""
        for key, value in defaults.items():
            session_key = f"{prefix}_{key}"
            if session_key not in st.session_state:
                st.session_state[session_key] = value
    
    def get_session_state(self, prefix: str, key: str, default: Any = None) -> Any:
        """Get a value from session state with the framework prefix"""
        session_key = f"{prefix}_{key}"
        return st.session_state.get(session_key, default)
    
    def set_session_state(self, prefix: str, key: str, value: Any) -> None:
        """Set a value in session state with the framework prefix"""
        session_key = f"{prefix}_{key}"
        st.session_state[session_key] = value
    
    def render_form_field(self, field_type: str, label: str, key: str, 
                          help_text: str = "", placeholder: str = "", 
                          options: List[Any] = None, **kwargs) -> Any:
        """Render a form field based on type and return the value"""
        if field_type == "text_input":
            return st.text_input(label=label, key=key, help=help_text, 
                                placeholder=placeholder, **kwargs)
        elif field_type == "text_area":
            return st.text_area(label=label, key=key, help=help_text, 
                               placeholder=placeholder, **kwargs)
        elif field_type == "selectbox":
            if not options:
                raise ValueError("Options must be provided for selectbox")
            return st.selectbox(label=label, options=options, key=key, 
                               help=help_text, **kwargs)
        elif field_type == "multiselect":
            if not options:
                raise ValueError("Options must be provided for multiselect")
            return st.multiselect(label=label, options=options, key=key, 
                                 help=help_text, **kwargs)
        elif field_type == "slider":
            return st.slider(label=label, key=key, help=help_text, **kwargs)
        elif field_type == "checkbox":
            return st.checkbox(label=label, key=key, help=help_text, **kwargs)
        elif field_type == "radio":
            if not options:
                raise ValueError("Options must be provided for radio")
            return st.radio(label=label, options=options, key=key, 
                           help=help_text, **kwargs)
        else:
            raise ValueError(f"Unsupported field type: {field_type}")
    
    def render_form(self, form_id: str, fields: List[Dict[str, Any]], 
                   submit_label: str = "Submit", 
                   on_submit: Optional[Callable] = None) -> Tuple[bool, Dict[str, Any]]:
        """Render a form with the specified fields and return submission status and values"""
        with st.form(form_id):
            values = {}
            for field in fields:
                field_type = field.get("type", "text_input")
                label = field.get("label", "")
                key = field.get("key", "")
                help_text = field.get("help", "")
                placeholder = field.get("placeholder", "")
                options = field.get("options", None)
                kwargs = {k: v for k, v in field.items() 
                         if k not in ["type", "label", "key", "help", "placeholder", "options"]}
                
                values[key] = self.render_form_field(
                    field_type, label, key, help_text, placeholder, options, **kwargs
                )
            
            submitted = st.form_submit_button(submit_label)
            
            if submitted and on_submit:
                on_submit(values)
                
            return submitted, values
    
    def validate_form_data(self, values: Dict[str, Any], 
                          required_fields: List[str]) -> Tuple[bool, List[str]]:
        """Validate form data and return validation status and error messages"""
        errors = []
        
        for field in required_fields:
            if not values.get(field):
                errors.append(f"Field '{field}' is required")
        
        return len(errors) == 0, errors
    
    def display_validation_errors(self, errors: List[str]) -> None:
        """Display validation errors in the UI"""
        for error in errors:
            st.error(error)
    
    def generate_ai_suggestion(self, prompt: str, model: str = "gpt-4o-mini") -> str:
        """Generate an AI suggestion using the specified prompt and model"""
        try:
            from utilities.gpt import get_completion
            return get_completion(prompt, model=model)
        except ImportError:
            st.warning("AI suggestions are not available. Please install the required packages.")
            return ""
        except Exception as e:
            st.error(f"Error generating AI suggestion: {e}")
            return ""
    
    def export_to_docx(self, title: str, sections: List[Dict[str, Any]]) -> Optional[bytes]:
        """Export framework data to a Word document"""
        try:
            from docx import Document
            from utilities.helpers import add_markdown_paragraph
            
            doc = Document()
            doc.add_heading(title, level=1)
            
            # Add metadata
            doc.add_heading("Analysis Information", level=2)
            doc.add_paragraph(f"Framework: {self.name}")
            doc.add_paragraph(f"Created: {self.metadata.get('created', 'Unknown')}")
            doc.add_paragraph(f"Last Modified: {self.metadata.get('last_modified', 'Unknown')}")
            
            # Add sections
            for section in sections:
                heading = section.get("heading", "")
                content = section.get("content", "")
                level = section.get("level", 2)
                
                if heading:
                    doc.add_heading(heading, level=level)
                
                if content:
                    add_markdown_paragraph(doc, content)
            
            # Save to bytes
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            return docx_bytes.getvalue()
        except ImportError:
            st.warning("Word export is not available. Please install python-docx.")
            return None
        except Exception as e:
            st.error(f"Error exporting to Word: {e}")
            return None
    
    def display_download_options(self, data: Dict[str, Any], 
                                filename_prefix: str) -> None:
        """Display download options for the framework data"""
        col1, col2 = st.columns(2)
        
        with col1:
            # JSON export
            json_data = self.export_to_json()
            st.download_button(
                label="Download JSON",
                data=json_data,
                file_name=f"{filename_prefix}_{self.name}.json",
                mime="application/json"
            )
        
        with col2:
            # Word export
            sections = [
                {"heading": "Overview", "content": f"Analysis using {self.name} framework", "level": 2},
                {"heading": "Components", "content": ", ".join(self.components), "level": 2}
            ]
            
            for component in self.components:
                response = self.get_response(component)
                if response:
                    sections.append({
                        "heading": component.capitalize(),
                        "content": response,
                        "level": 3
                    })
            
            docx_bytes = self.export_to_docx(f"{self.name} Analysis", sections)
            if docx_bytes:
                st.download_button(
                    label="Download Word Document",
                    data=docx_bytes,
                    file_name=f"{filename_prefix}_{self.name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )