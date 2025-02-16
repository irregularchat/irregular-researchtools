import streamlit as st
from docx import Document
from docx.shared import Pt
import re
import io
import openpyxl
from openpyxl.styles import NamedStyle, Font, PatternFill, Border, Side
import pandas as pd
from io import BytesIO

def initialize_session_keys(keys, default=""):
    for key in keys:
        if key not in st.session_state:
            st.session_state[key] = default


initialize_session_keys(["dotmlpf_summary", "tradoc_alignment", "command_endorsement", "improved_problem_statement"])

def remove_markdown(md_text: str) -> str:
    """Remove markdown formatting from text."""
    if not md_text:
        return ""
    # Remove bold, italics, inline code and links
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', md_text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'_(.*?)_', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'^#+\s', '', text, flags=re.MULTILINE)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text

def parse_markdown(text: str):
    """
    Parse a markdown string into a list of segments.
    
    Each segment is a tuple of (text, formatting_dict)
    where formatting_dict may contain keys like 'bold' or 'italic'.
    This is a very basic parser and supports **bold**, *italic*,
    __bold__, and _italic_ markers.
    """
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_)'
    segments = []
    pos = 0
    for m in re.finditer(pattern, text):
        start, end = m.span()
        if start > pos:
            segments.append((text[pos:start], {}))
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            segments.append((token[2:-2], {"bold": True}))
        elif token.startswith("*") and token.endswith("*"):
            segments.append((token[1:-1], {"italic": True}))
        elif token.startswith("__") and token.endswith("__"):
            segments.append((token[2:-2], {"bold": True}))
        elif token.startswith("_") and token.endswith("_"):
            segments.append((token[1:-1], {"italic": True}))
        pos = end
    if pos < len(text):
        segments.append((text[pos:], {}))
    return segments

def add_markdown_paragraph(document, markdown_text: str):
    """
    Add a paragraph to the DOCX document with markdown formatting.
    
    This function uses the basic parser (parse_markdown) to create
    runs that render **bold** and *italic* style without leaving markdown symbols.
    """
    paragraph = document.add_paragraph()
    segments = parse_markdown(markdown_text)
    for seg_text, fmt in segments:
        run = paragraph.add_run(seg_text)
        if fmt.get("bold"):
            run.bold = True
        if fmt.get("italic"):
            run.italic = True
    return paragraph

def add_section_content(document: Document, section_title: str, content: str):
    """
    Add content for a DOCX section. If the section is a Questions section,
    process each line so that the question text is rendered as a header (level 2)
    with bold and italic formatting and the search query is formatted in italic.
    
    For non-question sections, the content is rendered as a single markdown paragraph.
    """
    # If the section title indicates that it contains questions, process line-by-line.
    if "Questions" in section_title:
        lines = content.splitlines()
        for line in lines:
            line = line.strip().rstrip(";")
            if not line:
                continue
            # Look for the standard question format separated by '|'
            if '|' in line:
                parts = line.split('|', 1)
                if len(parts) == 2:
                    question_text = parts[0].strip()
                    search_text = parts[1].strip()
                    # Remove any existing prefixes.
                    if question_text.lower().startswith("question:"):
                        question_text = question_text[len("question:"):].strip()
                    if search_text.lower().startswith("search:"):
                        search_text = search_text[len("search:"):].strip()
                    # Add question as a header (level 2)
                    document.add_heading(question_text, level=2)
                    # Add associated search query in italic.
                    search_para = document.add_paragraph()
                    search_run = search_para.add_run("Search: " + search_text)
                    search_run.italic = True
                else:
                    add_markdown_paragraph(document, line)
            else:
                add_markdown_paragraph(document, line)
    else:
        add_markdown_paragraph(document, content)

def create_docx_document(title: str, sections: dict) -> io.BytesIO:
    """
    Create a DOCX document with a specified title and sections.
    
    This function now leverages our markdown rendering helpers to convert
    markdown formatting (bold, italics, headers for questions) to DOCX styling.
    
    Args:
      title: The document title.
      sections: A dictionary where keys are section headings and values are section contents (plain text).
      
    Returns:
      A BytesIO object containing the DOCX file.
    """
    document = Document()
    # Set default font
    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    document.add_heading(title, level=0)
    for heading, content in sections.items():
        document.add_heading(heading, level=1)
        add_section_content(document, heading, content)

    docx_io = io.BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    return docx_io

def apply_excel_styling(ws, headers: list, total_row: int):
    """
    Apply common Excel styling using openpyxl.
    
    This applies header styles, auto-adjusts column widths and sets thin borders.
    """
    # Create a header style
    header_style = NamedStyle(name='header')
    header_style.font = Font(bold=True)
    header_style.fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
    
    # Add the style if not already present
    if 'header' not in ws.parent.named_styles:
        ws.parent.add_named_style(header_style)

    # Set header cells style
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.style = header_style

    # Auto-adjust column widths
    for column_cells in ws.columns:
        max_length = 0
        for cell in column_cells:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column_cells[0].column_letter].width = max_length + 2

    # Apply borders to all cells in the range
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    for row in ws.iter_rows(min_row=1, max_row=total_row, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.border = thin_border

def export_ach_matrix_to_excel(hypotheses_list, evidence_list, ach_matrix, evidence_weights) -> bytes:
    """
    Export the ACH matrix to Excel using Pandas for basic tabular export.
    
    If detailed formatting is required, consider using openpyxl to further adjust styling.
    """
    data = []
    for ev in evidence_list:
        row = {"Evidence": ev, "Weight": evidence_weights.get(ev, 0)}
        for hyp in hypotheses_list:
            row[hyp] = ach_matrix.get((ev, hyp), "N/A")
        data.append(row)
    df = pd.DataFrame(data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='ACH Matrix')
        writer.save()
    output.seek(0)
    return output.getvalue()
